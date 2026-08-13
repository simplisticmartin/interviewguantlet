"""Resume/job-description text extraction.

Uploaded files are UNTRUSTED (spec section 44). This module extracts plain text and
nothing else: it does not follow links, does not execute embedded content, and does not
let document contents influence any control flow. Everything it returns is treated as
data and fenced before it reaches a model.
"""

from __future__ import annotations

import io
import re
import unicodedata

SUPPORTED_CONTENT_TYPES = {
    "text/plain": "txt",
    "text/markdown": "txt",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

MAX_EXTRACTED_CHARS = 60_000
# Zero-width and bidirectional control characters can hide instructions from a human
# reviewer while remaining visible to a model. Strip them at the boundary.
_INVISIBLE = re.compile(r"[​-‏‪-‮⁠-⁤﻿]")


class DocumentError(ValueError):
    """The upload could not be read as a document."""


def detect_kind(filename: str, content_type: str | None) -> str:
    if content_type and content_type.split(";")[0].strip() in SUPPORTED_CONTENT_TYPES:
        return SUPPORTED_CONTENT_TYPES[content_type.split(";")[0].strip()]
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return "pdf"
    if lowered.endswith(".docx"):
        return "docx"
    if lowered.endswith((".txt", ".md", ".markdown")):
        return "txt"
    raise DocumentError(
        f"Unsupported file type '{filename}'. Upload a PDF, DOCX, TXT, or Markdown file."
    )


def extract_text(data: bytes, filename: str, content_type: str | None = None) -> str:
    kind = detect_kind(filename, content_type)
    if kind == "pdf":
        text = _extract_pdf(data)
    elif kind == "docx":
        text = _extract_docx(data)
    else:
        text = _decode_text(data)

    cleaned = sanitise(text)
    if not cleaned.strip():
        raise DocumentError(
            "No readable text found. If this is a scanned PDF, paste the text instead."
        )
    return cleaned


def sanitise(text: str) -> str:
    """Normalise, strip invisible control characters, and bound the length."""
    normalised = unicodedata.normalize("NFKC", text)
    normalised = _INVISIBLE.sub("", normalised)
    normalised = normalised.replace("\r\n", "\n").replace("\r", "\n")
    normalised = re.sub(r"[ \t]+", " ", normalised)
    normalised = re.sub(r"\n{3,}", "\n\n", normalised)
    return normalised.strip()[:MAX_EXTRACTED_CHARS]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError("Could not decode the file as text.")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise DocumentError("Encrypted PDFs are not supported. Remove the password first.")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except DocumentError:
        raise
    except (PdfReadError, OSError, ValueError) as exc:
        raise DocumentError(f"Could not read the PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except (PackageNotFoundError, KeyError, ValueError) as exc:
        raise DocumentError(f"Could not read the DOCX file: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
